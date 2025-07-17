#!/usr/bin/env python
"""
Enhanced script to extract detailed content from ideal treatment sequence documents
"""

import os
import json
from docx import Document
from pathlib import Path

def extract_comprehensive_content(file_path):
    """Extract all content from a Word document including tables"""
    try:
        doc = Document(file_path)
        content_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(f"TEXT: {paragraph.text.strip()}")
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            content_parts.append(f"\nTABLE {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_cells.append(cell.text.strip())
                if row_cells:
                    content_parts.append(f"Row {row_idx + 1}: {' | '.join(row_cells)}")
        
        # Extract other document properties
        props = doc.core_properties
        metadata = {
            "title": props.title if props.title else "",
            "author": props.author if props.author else "",
            "subject": props.subject if props.subject else "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else ""
        }
        
        return {
            "content": "\n".join(content_parts),
            "metadata": metadata,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def main():
    """Extract comprehensive content from all ideal sequence documents"""
    sequences_dir = Path("ideal_sequences_of_treatements")
    
    if not sequences_dir.exists():
        print(f"Directory {sequences_dir} not found")
        return
    
    extracted_sequences = []
    
    # Process all .docx files in the directory
    for file_path in sequences_dir.glob("*.docx"):
        print(f"Processing {file_path}...")
        
        result = extract_comprehensive_content(file_path)
        if result:
            sequence_data = {
                "filename": file_path.name,
                "title": file_path.stem.replace("_", " ").title(),
                "content": result["content"],
                "metadata": result["metadata"],
                "stats": {
                    "paragraph_count": result["paragraph_count"],
                    "table_count": result["table_count"],
                    "content_length": len(result["content"])
                },
                "type": "ideal_sequence",
                "source": "dentist_guidelines"
            }
            extracted_sequences.append(sequence_data)
            print(f"Extracted {len(result['content'])} characters, {result['paragraph_count']} paragraphs, {result['table_count']} tables")
    
    # Save extracted sequences to JSON
    output_path = Path("DATA/IDEAL_SEQUENCES")
    output_path.mkdir(exist_ok=True)
    
    output_file = output_path / "detailed_ideal_sequences.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(extracted_sequences, f, indent=2, ensure_ascii=False)
    
    print(f"\nExtracted {len(extracted_sequences)} ideal sequences")
    print(f"Saved to {output_file}")
    
    return extracted_sequences

if __name__ == "__main__":
    sequences = main()
    
    # Display detailed preview of extracted content
    for sequence in sequences:
        print(f"\n{'='*50}")
        print(f"File: {sequence['filename']}")
        print(f"Title: {sequence['title']}")
        print(f"Stats: {sequence['stats']}")
        print(f"Metadata: {sequence['metadata']}")
        print(f"Content preview:\n{sequence['content'][:500]}...")
        print(f"{'='*50}")